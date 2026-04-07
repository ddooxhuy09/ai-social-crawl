from datetime import datetime
from pathlib import Path
import io
import base64
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse

from projects.db import _get_project, _update_project

router = APIRouter()

@router.post("/{project_id}/final/generate-word")
async def generate_word(project_id: str):
    """Generate Word document from redesign tasks."""
    project = _get_project(project_id)
    tasks = project["redesign"].get("tasks", [])
    done_tasks = [t for t in tasks if t.get("status") == "done"]
    if not done_tasks:
        raise HTTPException(status_code=400, detail="No completed redesign tasks found")

    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading(project.get("name", "Project"), 0)

        original = project.get("original", {})
        if original.get("original_item"):
            doc.add_heading("Original Product", level=1)
            item = original["original_item"]
            doc.add_paragraph(f"Name: {item.get('title', item.get('name', ''))}")
            doc.add_paragraph(f"URL: {item.get('url', '')}")

        doc.add_heading("Redesign Items", level=1)
        for i, task in enumerate(done_tasks, 1):
            doc.add_heading(f"Design {i}", level=2)
            src = task.get("source_item", {})
            doc.add_paragraph(f"Source: {src.get('title', src.get('url', ''))}")
            attrs = task.get("attributes", {})
            if attrs:
                doc.add_paragraph("Attributes:")
                for k, v in attrs.items():
                    if v:
                        doc.add_paragraph(f"  • {k}: {v}", style="List Bullet")
            img_data = task.get("generated_image")
            if img_data and img_data.startswith("data:image"):
                try:
                    b64 = img_data.split(",", 1)[1]
                    img_bytes = base64.b64decode(b64)
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(4))
                except Exception:
                    pass

        # Save file
        word_dir = Path("projects/word")
        word_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{project.get('name', 'project').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = word_dir / filename
        doc.save(str(filepath))

        project["final"]["status"] = "done"
        project["final"]["word_file"] = filename
        _update_project(project_id, project)
        
        try:
            from projects.telegram_notify import notify_step3_done
            notify_step3_done(project.get("name", "Unknown Project"), filename)
        except Exception as e:
            print(f"[TELEGRAM] Lỗi báo cáo Step 3: {e}")

        return FileResponse(
            path=str(filepath),
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word generation failed: {e}")
