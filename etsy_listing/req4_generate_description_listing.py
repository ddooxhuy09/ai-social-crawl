import asyncio
import io
import json
import re

from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel

from create_image_by_ai.image_generator import GEMINI_TEXT_MODEL, _get_gemini_client
from etsy_listing.shared import (
    build_listing_history_response,
    get_prompt_config,
    load_listing_history,
    now_iso,
    raise_for_ai_error,
    save_listing_history,
    strip_code_fence,
)


router = APIRouter()

KEYWORD_CONTEXT_LIMIT = 25


class GenerateDescriptionRequest(BaseModel):
    listing_name: str
    listing_title: str
    materials_skill_level: str = ""
    finished_sizes: str = ""
    story_ideas: str = ""
    shop_link: str = ""
    seed_keyword: str = ""


def _score_value(item: dict) -> float:
    try:
        return float(item.get("score") or 0)
    except (TypeError, ValueError):
        return 0.0


def _clean_text(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r"[\r\n\t]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip(" -*\u2022")


def _normalize_lines(value) -> list[str]:
    if isinstance(value, list):
        raw_lines = [str(item) for item in value]
    elif isinstance(value, str):
        raw_lines = re.split(r"[\r\n]+", value)
    else:
        raw_lines = []

    lines = []
    seen = set()
    for raw in raw_lines:
        cleaned = _clean_text(raw)
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        lines.append(cleaned)
    return lines


def _normalize_paragraph(value) -> str:
    if isinstance(value, list):
        parts = [_clean_text(str(item)) for item in value]
        return " ".join(part for part in parts if part)
    return _clean_text(str(value))


def _build_generation_prompt(
    req: GenerateDescriptionRequest,
    sorted_items: list[dict],
    req3_tags: list[str],
) -> str:
    prompt_obj = get_prompt_config("req4_description_generator")
    keyword_context = "\n".join(
        f"- {item.get('keyword', '')} (score: {_score_value(item):.2f})"
        for item in sorted_items[:KEYWORD_CONTEXT_LIMIT]
    ) or "- none"
    tag_context = "\n".join(f"- {tag}" for tag in req3_tags[:13]) or "- none"

    task = prompt_obj.get("task", "")
    task = task.replace("{seed_keyword}", req.seed_keyword)
    task = task.replace("{listing_title}", req.listing_title)
    task = task.replace("{materials_skill_level}", req.materials_skill_level or "None provided")
    task = task.replace("{finished_sizes}", req.finished_sizes or "None provided")
    task = task.replace("{story_ideas}", req.story_ideas or "None provided")
    task = task.replace("{shop_link}", req.shop_link or "None provided")
    task = task.replace("{ai_keywords}", keyword_context)
    task = task.replace("{req3_tags}", tag_context)

    rules = "\n".join(f"- {rule}" for rule in prompt_obj.get("rules", []))
    return f"Role: {prompt_obj.get('role', '')}\n\nTask: {task}\n\nRules:\n{rules}"


async def _generate_description_sections(
    req: GenerateDescriptionRequest,
    sorted_items: list[dict],
    req3_tags: list[str],
) -> dict:
    prompt = _build_generation_prompt(req, sorted_items, req3_tags)
    client = _get_gemini_client()
    response = await asyncio.to_thread(
        client.models.generate_content,
        model=GEMINI_TEXT_MODEL,
        contents=prompt,
        config={"temperature": 0.65},
    )
    text = strip_code_fence(response.text)
    parsed = json.loads(text)
    if not isinstance(parsed, dict):
        raise ValueError("REQ4 AI response was not a valid object.")
    return parsed


def _build_description_text(
    listing_title: str,
    download_includes_lines: list[str],
    materials_skill_level_lines: list[str],
    finished_size_lines: list[str],
    what_is_it_use_for_lines: list[str],
    behind_the_design_paragraph: str,
    please_read_before_purchase_lines: list[str],
    faqs_lines: list[str],
    shop_link: str,
) -> str:
    shop_link_line = shop_link.strip() or "https://www.etsy.com/shop/KniriCrochetHome"

    parts = [
        listing_title.strip(),
        "",
        "YOUR DOWNLOAD INCLUDES",
        "\n".join(f"• {line}" for line in download_includes_lines),
        "",
        "MATERIALS & SKILL LEVEL",
        "\n".join(f"• {line}" if ':' in line else line for line in materials_skill_level_lines),
        "",
        "FINISHED SIZE",
        "\n".join(finished_size_lines),
        "",
        "WHAT IS IT USE FOR",
        "\n".join(f"• {line}" for line in what_is_it_use_for_lines),
        "",
        "BEHIND THE DESIGN",
        behind_the_design_paragraph,
        "",
        "PLEASE READ BEFORE PURCHASE",
        "\n".join(f"• {line}" for line in please_read_before_purchase_lines),
        "",
        "FAQs",
        "\n".join(faqs_lines),
        "",
        "View more Kniri patterns:",
        shop_link_line,
    ]

    return "\n".join(part for part in parts).strip()


@router.post("/api/listing/generate_description")
async def generate_listing_description(req: GenerateDescriptionRequest):
    """
    REQ 4: Generate a structured Etsy listing description and save it into the same listing history.
    """
    history = load_listing_history(req.listing_name)
    if not history or not history.get("req1"):
        raise HTTPException(status_code=404, detail="Không tìm thấy dữ liệu REQ1 cho listing này.")
    if not req.listing_title.strip():
        raise HTTPException(status_code=400, detail="Vui lòng nhập title để tạo description.")

    items = history["req1"].get("data") or []
    if not items:
        raise HTTPException(status_code=400, detail="History REQ1 không có keyword để tạo description.")

    seed_keyword = req.seed_keyword or history.get("seed_keyword") or ""
    req3_tags = ((history.get("req3") or {}).get("tags")) or []
    sorted_items = sorted(items, key=_score_value, reverse=True)
    keyword_sources = [str(item.get("keyword") or "").strip() for item in sorted_items[:12] if str(item.get("keyword") or "").strip()]

    work_req = GenerateDescriptionRequest(
        listing_name=req.listing_name,
        listing_title=req.listing_title.strip(),
        materials_skill_level=req.materials_skill_level.strip(),
        finished_sizes=req.finished_sizes.strip(),
        story_ideas=req.story_ideas.strip(),
        shop_link=req.shop_link.strip(),
        seed_keyword=seed_keyword,
    )

    try:
        generated = await _generate_description_sections(work_req, sorted_items, req3_tags)
        download_includes_lines = _normalize_lines(generated.get("download_includes_lines"))
        materials_skill_level_lines = _normalize_lines(generated.get("materials_skill_level_lines"))
        finished_size_lines = _normalize_lines(generated.get("finished_size_lines"))
        what_is_it_use_for_lines = _normalize_lines(generated.get("what_is_it_use_for_lines"))
        behind_the_design_paragraph = _normalize_paragraph(generated.get("behind_the_design_paragraph"))
        please_read_before_purchase_lines = _normalize_lines(generated.get("please_read_before_purchase_lines"))
        faqs_lines = _normalize_lines(generated.get("faqs_lines"))

        if not behind_the_design_paragraph:
            raise ValueError("REQ4 missing behind_the_design_paragraph.")
    except Exception as e:
        raise_for_ai_error(e, "REQ4 generate description")

    description_text = _build_description_text(
        listing_title=work_req.listing_title,
        download_includes_lines=download_includes_lines,
        materials_skill_level_lines=materials_skill_level_lines,
        finished_size_lines=finished_size_lines,
        what_is_it_use_for_lines=what_is_it_use_for_lines,
        behind_the_design_paragraph=behind_the_design_paragraph,
        please_read_before_purchase_lines=please_read_before_purchase_lines,
        faqs_lines=faqs_lines,
        shop_link=work_req.shop_link,
    )

    history["seed_keyword"] = seed_keyword
    history["req4"] = {
        "listing_title": work_req.listing_title,
        "materials_skill_level": work_req.materials_skill_level,
        "finished_sizes": work_req.finished_sizes,
        "story_ideas": work_req.story_ideas,
        "shop_link": work_req.shop_link,
        "keyword_sources": keyword_sources,
        "sections": {
            "download_includes_lines": download_includes_lines,
            "materials_skill_level_lines": materials_skill_level_lines,
            "finished_size_lines": finished_size_lines,
            "what_is_it_use_for_lines": what_is_it_use_for_lines,
            "behind_the_design_paragraph": behind_the_design_paragraph,
            "please_read_before_purchase_lines": please_read_before_purchase_lines,
            "faqs_lines": faqs_lines,
        },
        "description_text": description_text,
        "updated_at": now_iso(),
    }
    history = save_listing_history(history)
    return build_listing_history_response(history)


def _extract_text_from_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_text_from_pdf(content: bytes) -> str:
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(content))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n".join(parts)


CHUNK_SIZE = 6000
CHUNK_OVERLAP = 300


def _split_into_chunks(text: str) -> list[str]:
    """Split text into overlapping chunks, breaking at newlines where possible."""
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + CHUNK_SIZE, length)
        # Try to break at a newline boundary to avoid cutting mid-sentence
        if end < length:
            boundary = text.rfind("\n", start, end)
            if boundary > start:
                end = boundary + 1
        chunks.append(text[start:end])
        start = end - CHUNK_OVERLAP if end < length else length
    return chunks


def _build_extraction_prompt(text_chunk: str) -> str:
    prompt_obj = get_prompt_config("req4_pattern_extractor")
    task = prompt_obj.get("task", "").replace("{text_chunk}", text_chunk)
    rules = "\n".join(f"- {rule}" for rule in prompt_obj.get("rules", []))
    return f"Role: {prompt_obj.get('role', '')}\n\nTask: {task}\n\nRules:\n{rules}"


async def _extract_chunk(text_chunk: str) -> dict:
    prompt = _build_extraction_prompt(text_chunk)
    client = _get_gemini_client()
    response = await asyncio.to_thread(
        client.models.generate_content,
        model=GEMINI_TEXT_MODEL,
        contents=prompt,
        config={"temperature": 0.2},
    )
    raw = strip_code_fence(response.text or "")
    return json.loads(raw)


def _merge_chunk_results(results: list[dict]) -> dict:
    """Merge results from multiple chunks into one final dict.

    - listing_title: first non-empty value wins (title is usually in the first chunk)
    - other fields: concatenate unique non-empty values from all chunks
    """
    merged: dict[str, str] = {
        "listing_title": "",
        "materials_skill_level": "",
        "finished_sizes": "",
        "story_ideas": "",
    }
    for key in ("materials_skill_level", "finished_sizes", "story_ideas"):
        seen: list[str] = []
        for r in results:
            val = (r.get(key) or "").strip()
            if val and val not in seen:
                seen.append(val)
        merged[key] = "; ".join(seen)

    for r in results:
        val = (r.get("listing_title") or "").strip()
        if val:
            merged["listing_title"] = val
            break

    return merged


async def _extract_fields_with_gemini(text: str) -> dict:
    chunks = _split_into_chunks(text)
    results = await asyncio.gather(*[_extract_chunk(chunk) for chunk in chunks], return_exceptions=True)
    valid = [r for r in results if isinstance(r, dict)]
    if not valid:
        raise ValueError("All chunks failed AI extraction.")
    return _merge_chunk_results(valid)


@router.post("/api/listing/extract-pattern-info")
async def extract_pattern_info(file: UploadFile = File(...)):
    """Extract REQ4 fields (title, materials, sizes, story) from a .docx or .pdf pattern file."""
    filename = (file.filename or "").lower()
    if not (filename.endswith(".docx") or filename.endswith(".pdf")):
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are supported.")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    try:
        if filename.endswith(".docx"):
            text = _extract_text_from_docx(content)
        else:
            text = _extract_text_from_pdf(content)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not read file: {e}")

    if not text.strip():
        raise HTTPException(status_code=422, detail="No text could be extracted from the file.")

    try:
        fields = await _extract_fields_with_gemini(text)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI extraction failed: {e}")

    return {
        "listing_title": fields.get("listing_title", ""),
        "materials_skill_level": fields.get("materials_skill_level", ""),
        "finished_sizes": fields.get("finished_sizes", ""),
        "story_ideas": fields.get("story_ideas", ""),
    }
